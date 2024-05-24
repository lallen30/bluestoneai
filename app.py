import uuid
from flask import Flask, request, jsonify, g
from flask_cors import CORS
from flask_jwt_extended import JWTManager
from functools import wraps
import argparse
from langchain.schema.document import Document
from langchain_community.vectorstores import Chroma
from langchain.prompts import ChatPromptTemplate
import openai
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import os
import sqlite3
import logging
import traceback
import hashlib
import pandas as pd
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import requests
from bs4 import BeautifulSoup
from populate_database import main as populate_main
from get_embedding_function import get_embedding_function

logging.basicConfig(level=logging.INFO)

app = Flask(__name__)
CORS(app, supports_credentials=True)
jwt = JWTManager(app)
load_dotenv()
app.secret_key = os.getenv('SECRET_KEY')
DATABASE = 'qna.db'
DATA_FOLDER = './data'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'csv', 'doc', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

CHROMA_PATH = "chroma"

PROMPT_TEMPLATE = """
Answer the question based only on the following context:

{context}

---

Answer the question based on the above context: {question}
"""

openai.api_key = os.getenv('OPENAI_API_KEY')

def get_chroma_instance():
    return Chroma(persist_directory=CHROMA_PATH, embedding_function=get_embedding_function())

@app.after_request
def after_request_func(response):
    origin = request.headers.get('Origin')
    if origin:
        response.headers.add('Access-Control-Allow-Origin', origin)
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        response.headers.add('Access-Control-Allow-Methods', 'GET,POST')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    return response

def get_db():
    db = getattr(g, '_database', None)
    if db is None:
        db = g._database = sqlite3.connect(DATABASE)
    return db

@app.teardown_appcontext
def close_connection(exception):
    db = getattr(g, '_database', None)
    if db is not None:
        db.close()

def init_db():
    with app.app_context():
        db = get_db()
        cursor = db.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS system_input_text (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                text_type TEXT,
                text TEXT,
                UNIQUE(text_type)
            )
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS api_keys (
                key_name TEXT PRIMARY KEY,
                key_value TEXT
            );
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS user_questions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                question TEXT,
                answer TEXT,
                question_session_id TEXT,
                asked_on DATETIME DEFAULT CURRENT_TIMESTAMP
            );
        """)
        db.commit()

init_db()

def require_api_key(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        auth_header = request.headers.get('Authorization')
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header.split(' ')[1]
            if token == os.getenv('STATIC_API_TOKEN'):
                return f(*args, **kwargs)
        return jsonify({"msg": "Invalid API key"}), 401
    return decorated_function

def get_openai_api_key():
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT key_value FROM api_keys WHERE key_name = 'OPENAI_API_KEY'")
    row = cursor.fetchone()
    return row[0] if row else None

@app.route('/set_api_key', methods=['POST'])
@require_api_key
def set_api_key():
    try:
        content = request.json
        key_name = content.get('key_name')
        key_value = content.get('key_value')

        if not key_name or not key_value:
            return jsonify({'error': 'Missing key name or key value'}), 400

        db = get_db()
        cursor = db.cursor()
        cursor.execute("""
            INSERT INTO api_keys (key_name, key_value)
            VALUES (?, ?)
            ON CONFLICT(key_name)
            DO UPDATE SET key_value = excluded.key_value;
        """, (key_name, key_value))
        db.commit()
        os.environ['OPENAI_API_KEY'] = key_value
        cursor.execute("SELECT key_value FROM api_keys WHERE key_name = ?", (key_name,))
        updated_key_row = cursor.fetchone()
        updated_key = updated_key_row[0] if updated_key_row else None

        return jsonify({'message': 'API key updated successfully.', 'updated_key': updated_key})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/populate_db', methods=['POST'])
@require_api_key
def populate_db():
    try:
        reset = request.json.get('reset', False)
        populate_main(reset=reset)
        return jsonify({'message': 'Database population initiated successfully.'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/submit_text', methods=['POST'])
@require_api_key
def submit_texts():
    contents = request.json

    try:
        db = get_db()
        cursor = db.cursor()

        for content in contents:
            text_type = content.get('text_type')
            text = content.get('text')

            if not text or not text_type:
                return jsonify({'error': 'No text or text type provided for one of the entries'}), 400

            cursor.execute("""
                INSERT INTO system_input_text (text_type, text)
                VALUES (?, ?)
                ON CONFLICT(text_type)
                DO UPDATE SET text = excluded.text;
            """, (text_type, text))

        db.commit()
        return jsonify({'message': 'Texts updated successfully.'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/qna', methods=['POST'])
@require_api_key
def qna():
    question = request.json.get('question')
    question_session_id = request.json.get('question_session_id', str(uuid.uuid4()))

    if not question:
        return jsonify({'error': 'No question provided.'}), 400

    try:
        db = get_db()
        cursor = db.cursor()

        # Insert the current question and session ID into the database
        cursor.execute("""
            INSERT INTO user_questions (question, question_session_id)
            VALUES (?, ?)
        """, (question, question_session_id))
        question_id = cursor.lastrowid
        db.commit()

        # Fetch all prior Q&As in this session for context
        cursor.execute("""
            SELECT question, answer FROM user_questions
            WHERE question_session_id = ? AND id < ?
            ORDER BY asked_on DESC
        """, (question_session_id, question_id))
        previous_qas = cursor.fetchall()

        # Retrieve intro_message and system_prompt from the database
        intro_message = get_system_input_text('intro_message')
        system_prompt = get_system_input_text('system_prompt')

        # Construct context for the current question
        context = ' '.join([f"Q: {qa[0]} A: {qa[1]}" for qa in previous_qas if qa[1]])
        full_query = f"{intro_message}\n\n{context} Current Question: {question}"

        # Invoke the model
        response_text, sources = query_rag(full_query, system_prompt)

        # Update the database with the answer
        cursor.execute("""
            UPDATE user_questions
            SET answer = ?
            WHERE id = ?
        """, (response_text, question_id))
        db.commit()

        return jsonify({
            'question_session_id': question_session_id,
            'question': question,
            'answer': response_text,
            'sources': sources
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def get_system_input_text(text_type):
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT text FROM system_input_text WHERE text_type = ?", (text_type,))
    row = cursor.fetchone()
    return row[0] if row else ""


def query_rag(query_text: str, system_prompt: str):
    db = get_chroma_instance()
    results = db.similarity_search_with_score(query_text, k=5)

    if not results:  # Check if no results are returned
        return "No relevant data found to answer the question.", []

    context_text = "\n\n---\n\n".join([doc.page_content for doc, _score in results])
    prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
    prompt = prompt_template.format(context=context_text, question=query_text)

    # Combine system_prompt with the formatted prompt
    full_prompt = f"{system_prompt}\n\n{prompt}"

    response = openai.chat.completions.create(
        # model="gpt-3.5-turbo",
        model="gpt-4o",
        messages=[
            {"role": "user", "content": full_prompt}
        ]
    )

    response_text = response.choices[0].message.content.strip()
    if response_text.startswith("A:"):
        response_text = response_text[2:].strip()

    sources = [doc.metadata.get("source", None) for doc, _score in results]
    return response_text, sources



def calculate_chunk_ids(chunks):
    last_page_id = None
    current_chunk_index = 0

    for chunk in chunks:
        source = chunk.metadata.get("source")
        page = chunk.metadata.get("page")
        current_page_id = f"{source}:{page}"

        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        chunk_id = f"{current_page_id}:{current_chunk_index}"
        last_page_id = current_page_id
        chunk.metadata["id"] = chunk_id

    return chunks

def vectorize_and_store_file(file_path):
    db = get_chroma_instance()
    collection = db._collection
    chunks = []

    try:
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                logging.info(f"Number of pages in PDF: {len(pdf_reader.pages)}")

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    logging.info(f"Extracted text from page {page_num}: {text[:100]}")

                    chunk = Document(page_content=text, metadata={"source": file_path, "page": page_num})
                    chunks.append(chunk)
        
        elif file_path.endswith('.docx'):
            docx_document = DocxDocument(file_path)
            full_text = []
            for para in docx_document.paragraphs:
                full_text.append(para.text)
            text = '\n'.join(full_text)
            logging.info(f"Extracted text from DOCX: {text[:100]}")

            chunk = Document(page_content=text, metadata={"source": file_path})
            chunks.append(chunk)

        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                logging.info(f"Extracted text from TXT: {text[:100]}")

                chunk = Document(page_content=text, metadata={"source": file_path})
                chunks.append(chunk)

        elif file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
            text = df.to_string()
            logging.info(f"Extracted text from CSV: {text[:100]}")

            chunk = Document(page_content=text, metadata={"source": file_path})
            chunks.append(chunk)

        else:
            raise ValueError("Unsupported file type")

        chunks_with_ids = calculate_chunk_ids(chunks)
        for chunk in chunks_with_ids:
            text = chunk.page_content
            embedding_function = get_embedding_function()
            embeddings = embedding_function.embed_documents([text])
            logging.info(f"Embeddings for page {chunk.metadata.get('page', 'n/a')}: {embeddings}")

            collection.add(
                documents=[text], 
                metadatas=[chunk.metadata], 
                embeddings=embeddings,
                ids=[chunk.metadata['id']]
            )
    except Exception as e:
        logging.error(f"Error vectorizing file {file_path}: {e}")
        logging.error(traceback.format_exc())
        raise

@app.route('/remove_vectorized_data', methods=['DELETE'])
@require_api_key
def delete_file():
    if not request.json or 'filename' not in request.json:
        return jsonify({'error': 'Filename is required in JSON payload'}), 400

    filename = request.json['filename']
    safe_filename = secure_filename(filename)
    file_path = os.path.join('./data', safe_filename)

    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404

    try:
        os.remove(file_path)
        logging.info(f"File {file_path} deleted from the filesystem.")
        remove_data_from_chroma(safe_filename)
        return jsonify({'message': 'File and associated data deleted successfully'}), 200
    except Exception as e:
        logging.error(f"Error while deleting file {filename}: {e}")
        return jsonify({'error': str(e)}), 500

def fetch_all_chunk_ids():
    db = get_chroma_instance()
    collection = db._collection

    # Retrieve all documents to extract their IDs
    result = collection.get(include=["documents", "metadatas"])
    ids = result.get("ids", [])  # IDs are always included by default

    return ids

def remove_data_from_chroma(filename):
    db = get_chroma_instance()
    collection = db._collection

    # Retrieve all chunk IDs
    chunk_ids = fetch_all_chunk_ids()
    
    # Filter chunk IDs that belong to the given filename
    ids_to_delete = [chunk_id for chunk_id in chunk_ids if filename in chunk_id]
    logging.info(f"IDs to delete for filename {filename}: {ids_to_delete}")

    if ids_to_delete:
        try:
            collection.delete(ids=ids_to_delete)
            logging.info(f"Deleted IDs for filename {filename}: {ids_to_delete}")
        except Exception as e:
            logging.error(f"Error while deleting IDs {ids_to_delete} for filename {filename}: {e}")
    else:
        logging.info(f"No IDs found for filename {filename}")

@app.route('/upload', methods=['POST'])
@require_api_key
def upload_file():
    if 'file' not in request.files:
        return jsonify({'message': 'No file part in the request', 'status': 'error'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'message': 'No selected file', 'status': 'error'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(DATA_FOLDER, filename)

        try:
            # Save the uploaded file to the designated folder
            file.save(file_path)

            # Process the file (vectorize and add to Chroma)
            vectorize_and_store_file(file_path)

            return jsonify({'message': 'File uploaded and processed successfully', 'status': 'success'}), 200
        except Exception as e:
            # Log the error with detailed stack trace
            logging.error(f"Error processing file {filename}: {e}")
            logging.error(traceback.format_exc())
            return jsonify({'message': 'Error processing file', 'status': 'error'}), 500
    else:
        return jsonify({'message': 'Invalid file type', 'status': 'error'}), 400

@app.route('/process_webpage', methods=['POST'])
@require_api_key
def process_webpage():
    content = request.json
    url = content.get('url')

    if not url:
        return jsonify({'error': 'No URL provided.'}), 400

    result = process_single_webpage_content(url)
    return jsonify(result)

def process_single_webpage_content(url):
    """
    Process content from a single URL by fetching the webpage content,
    saving it as a text file, and vectorizing and building the model.
    """
    try:
        # print(f"Processing URL: {url}")

        # Fetch the webpage content
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        text_content = soup.get_text()

        # Generate a hash of the URL to use as a filename
        url_hash = hashlib.md5(url.encode()).hexdigest()
        filename = f"{url_hash}.txt"

        # Save the webpage content as a text file
        save_path = os.path.join('data', filename)
        with open(save_path, 'w') as file:
            file.write(text_content)

        # print(f"Webpage content saved at {save_path}")

        # Vectorize and build the model with the webpage content
        vectorize_and_store_file(save_path)

        # print(f"Data for {url} processed and stored.")
        return {'status': 'success', 'message': f'Webpage content from {url} processed successfully.'}
    except Exception as e:
        # print(f"An error occurred while processing {url}: {e}")
        return {'status': 'error', 'message': str(e)}

@app.route('/get_all_chunk_ids', methods=['GET'])
@require_api_key
def get_all_chunk_ids():
    try:
        chunk_ids = fetch_all_chunk_ids()
        return jsonify({'chunk_ids': chunk_ids})
    except Exception as e:
        logging.error(f"Error while fetching all chunk IDs: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/test_connection', methods=['GET'])
def test():
    return jsonify({'message': 'Success'}), 200

@app.errorhandler(404)
def page_not_found(e):
    return jsonify({'error': 'Page not found'}), 404

@app.route('/')
def home():
    return jsonify({'message': 'Welcome to the Bluestone AI API!'})

if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=5001)
